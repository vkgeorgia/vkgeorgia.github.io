import os
import logging
import re
import io
from pathlib import Path
from typing import Dict, List, Optional
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload

try:
    import google.generativeai as genai
except ImportError:
    genai = None

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    from docx import Document
except ImportError:
    Document = None

logger = logging.getLogger(__name__)

class RAGService:
    def __init__(self):
        self.knowledge_base = []
        self.model = None
        self.drive_service = None
        self.drive_folder_id = None
        self._vector_search_available = False

        # Gemini Setup (Old SDK)
        api_key = os.getenv("GEMINI_API_KEY")
        if api_key and genai:
            try:
                genai.configure(api_key=api_key)
                self.model = genai.GenerativeModel('gemini-2.5-flash')
                logger.info(f"Gemini API configured successfully with gemini-2.5-flash")
            except Exception as e:
                logger.error(f"Failed to configure Gemini: {e}")
                self.model = None
        else:
            logger.warning("GEMINI_API_KEY not found or google.generativeai not installed.")
            self.model = None

        # Check if vector search is available (pgvector + knowledge_chunks populated)
        self._check_vector_search()

        # Load knowledge base
        self.load_knowledge_base()
    
    def _init_drive_service(self):
        """Initialize Google Drive API service."""
        try:
            credentials_path = Path(__file__).resolve().parent.parent.parent.parent / "google_credentials.json"
            if not credentials_path.exists():
                logger.warning("google_credentials.json not found. Drive integration disabled.")
                return
            
            SCOPES = ['https://www.googleapis.com/auth/drive.readonly']
            credentials = service_account.Credentials.from_service_account_file(
                str(credentials_path), scopes=SCOPES)
            self.drive_service = build('drive', 'v3', credentials=credentials)
            logger.info("Google Drive API initialized successfully")
        except Exception as e:
            logger.error(f"Failed to initialize Drive API: {e}")
            self.drive_service = None
    
    def load_knowledge_base(self):
        """Load documents from local knowledge base directories and DB."""
        logger.info("Loading knowledge base...")

        self._load_from_local()
        db_docs = self._load_from_db()
        self.knowledge_base.extend(db_docs)
        total_chars = sum(len(doc) for doc in self.knowledge_base)
        logger.info(f"Loaded {len(self.knowledge_base)} documents total. Total characters: {total_chars}")
    
    def _load_from_drive(self):
        """Load files from Google Drive folder."""
        query = f"'{self.drive_folder_id}' in parents and trashed=false"
        results = self.drive_service.files().list(
            q=query,
            fields="files(id, name, mimeType)",
            pageSize=100
        ).execute()
        
        files = results.get('files', [])
        logger.info(f"Found {len(files)} files in Drive folder")
        
        for file in files:
            try:
                content = self._download_and_parse_drive_file(file)
                if content:
                    self.knowledge_base.append(content)
            except Exception as e:
                logger.error(f"Failed to process {file['name']}: {e}")
    
    def _download_and_parse_drive_file(self, file):
        """Download and parse a single Drive file."""
        file_id = file['id']
        file_name = file['name']
        mime_type = file['mimeType']
        
        # Handle Google Docs specially
        if mime_type == 'application/vnd.google-apps.document':
            request = self.drive_service.files().export_media(fileId=file_id, mimeType='text/plain')
            text_buffer = io.BytesIO()
            downloader = MediaIoBaseDownload(text_buffer, request)
            done = False
            while not done:
                _, done = downloader.next_chunk()
            text_buffer.seek(0)
            return text_buffer.read().decode('utf-8', errors='ignore')
        
        # Download regular files
        request = self.drive_service.files().get_media(fileId=file_id)
        file_buffer = io.BytesIO()
        downloader = MediaIoBaseDownload(file_buffer, request)
        
        done = False
        while not done:
            status, done = downloader.next_chunk()
        
        file_buffer.seek(0)
        
        # Parse based on type
        if mime_type == 'text/plain' or file_name.endswith('.txt'):
            return file_buffer.read().decode('utf-8', errors='ignore')
        elif mime_type == 'application/pdf' or file_name.endswith('.pdf'):
            return self._parse_pdf_buffer(file_buffer)
        elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' or file_name.endswith('.docx'):
            return self._parse_docx_buffer(file_buffer)
        elif mime_type.startswith('application/vnd.google-apps'):
            logger.warning(f"Unsupported Google Apps file type: {mime_type} for {file_name}")
            return None
        else:
            logger.warning(f"Unsupported file type: {mime_type} for {file_name}")
            return None
    
    def _parse_pdf_buffer(self, buffer):
        """Parse PDF from buffer."""
        if not pdfplumber:
            logger.warning("pdfplumber not installed, skipping PDF")
            return None
        text = ""
        with pdfplumber.open(buffer) as pdf:
            for page in pdf.pages:
                text += page.extract_text() or ""
        return text
    
    def _parse_docx_buffer(self, buffer):
        """Parse DOCX from buffer."""
        if not Document:
            logger.warning("python-docx not installed, skipping DOCX")
            return None
        doc = Document(buffer)
        return "\n".join([para.text for para in doc.paragraphs])
    
    def _load_from_local(self):
        """Fallback: Load from local knowledge base directories."""
        candidate_dirs = []
        env_dir = os.getenv("KNOWLEDGE_BASE_DIR")
        if env_dir:
            candidate_dirs.append(Path(env_dir))

        # /app/knowledge_base when running in container
        candidate_dirs.append(Path(__file__).resolve().parent.parent.parent / "knowledge_base")
        # valerii-site repo root knowledge-base
        candidate_dirs.append(Path(__file__).resolve().parent.parent.parent.parent.parent / "knowledge-base")

        def load_from_dir(kb_dir: Path):
            # Read .md and .txt files (excluding sources/ directories)
            for pattern in ["*.md", "*.txt"]:
                for file in kb_dir.rglob(pattern):
                    if "sources" in file.parts:
                        continue
                    try:
                        with open(file, "r", encoding="utf-8") as f:
                            self.knowledge_base.append(f.read())
                    except Exception as e:
                        logger.error(f"Error reading {file}: {e}")

            # Read .pdf files (excluding sources/ directories)
            if pdfplumber:
                for pdf_file in kb_dir.rglob("*.pdf"):
                    if "sources" in pdf_file.parts:
                        continue
                    try:
                        with pdfplumber.open(pdf_file) as pdf:
                            text = ""
                            for page in pdf.pages:
                                text += page.extract_text() or ""
                            self.knowledge_base.append(text)
                    except Exception as e:
                        logger.error(f"Error reading {pdf_file}: {e}")

            # Read .docx files (excluding sources/ directories)
            if Document:
                for docx_file in kb_dir.rglob("*.docx"):
                    if "sources" in docx_file.parts:
                        continue
                    try:
                        doc = Document(docx_file)
                        text = "\n".join([para.text for para in doc.paragraphs])
                        self.knowledge_base.append(text)
                    except Exception as e:
                        logger.error(f"Error reading {docx_file}: {e}")

        seen = set()
        for kb_dir in candidate_dirs:
            # 1. Load from directory
            kb_dir = kb_dir.expanduser()
            kb_dir_key = str(kb_dir)
            if kb_dir_key in seen:
                continue
            seen.add(kb_dir_key)
            
            if not kb_dir.exists():
                logger.warning(f"Knowledge base directory not found: {kb_dir}")
            else:
                load_from_dir(kb_dir)
                
            # 2. ALSO Load specific root files from the parent repo if we are in the valerii-site structure
            # kb_dir usually points to .../knowledge-base. The parent of that is the repo root.
            repo_root = kb_dir.parent
            for root_file in ["profile.md", "business-challenges.md", "index.md"]:
                file_path = repo_root / root_file
                if file_path.exists() and file_path.is_file():
                    try:
                        with open(file_path, "r", encoding="utf-8") as f:
                            content = f.read()
                            # Prepend filename context so RAG knows what this is
                            self.knowledge_base.append(f"Content from {root_file}:\n{content}")
                            logger.info(f"Loaded root document: {root_file}")
                    except Exception as e:
                        logger.error(f"Error reading root file {file_path}: {e}")

    def _load_from_db(self) -> list:
        """Load all projects from Neon DB and format as rich-text documents."""
        db_url = os.getenv("NEON_DATABASE_URL")
        if not db_url:
            logger.warning("NEON_DATABASE_URL not set — skipping DB knowledge load")
            return []

        try:
            import psycopg  # psycopg3
        except ImportError:
            try:
                import psycopg2 as psycopg  # type: ignore
            except ImportError:
                logger.warning("psycopg / psycopg2 not installed — skipping DB knowledge load")
                return []

        docs = []
        try:
            conn = psycopg.connect(db_url)
            cur = conn.cursor()

            # Load all projects
            cur.execute("""
                SELECT
                    p.id, p.code, p.title, p.employer, p.client,
                    p.industry, p.domain, p.role,
                    p.year_start, p.year_end,
                    p.technology,
                    p.star_situation, p.star_task, p.star_action, p.star_result,
                    p.key_result, p.executive_md
                FROM projects p
                ORDER BY p.year_start DESC NULLS LAST, p.id
            """)
            cols = [d[0] for d in cur.description]
            rows = cur.fetchall()

            # Load domain/industry/role tags per project
            cur.execute("SELECT project_id, slug FROM project_domains")
            domain_map: dict = {}
            for pid, slug in cur.fetchall():
                domain_map.setdefault(pid, []).append(slug)

            cur.execute("SELECT project_id, slug FROM project_industries")
            industry_map: dict = {}
            for pid, slug in cur.fetchall():
                industry_map.setdefault(pid, []).append(slug)

            cur.execute("SELECT project_id, slug FROM project_roles")
            role_map: dict = {}
            for pid, slug in cur.fetchall():
                role_map.setdefault(pid, []).append(slug)

            conn.close()

            for row in rows:
                p = dict(zip(cols, row))
                pid = p["id"]
                parts = []
                period = f"{p['year_start'] or '?'}"
                if p.get("year_end") and p["year_end"] != p["year_start"]:
                    period += f"–{p['year_end']}"
                parts.append(f"## Project: {p['title']} ({p['code']})")
                parts.append(f"Period: {period} | Role: {p.get('role') or '—'}")
                parts.append(f"Employer: {p.get('employer') or '—'} | Client: {p.get('client') or '—'}")
                parts.append(f"Industry: {p.get('industry') or '—'} | Domain: {p.get('domain') or '—'}")
                if domain_map.get(pid):
                    parts.append(f"Domain tags: {', '.join(domain_map[pid])}")
                if industry_map.get(pid):
                    parts.append(f"Industry tags: {', '.join(industry_map[pid])}")
                if role_map.get(pid):
                    parts.append(f"Role tags: {', '.join(role_map[pid])}")
                if p.get("technology"):
                    parts.append(f"Technologies: {p['technology']}")
                if p.get("star_situation"):
                    parts.append(f"\nSituation: {p['star_situation']}")
                if p.get("star_task"):
                    parts.append(f"Task: {p['star_task']}")
                if p.get("star_action"):
                    parts.append(f"Action: {p['star_action']}")
                if p.get("star_result"):
                    parts.append(f"Result: {p['star_result']}")
                if p.get("key_result"):
                    parts.append(f"Key Result: {p['key_result']}")
                if p.get("executive_md"):
                    parts.append(f"\n{p['executive_md']}")
                docs.append("\n".join(parts))

            logger.info(f"Loaded {len(docs)} projects from DB")
        except Exception as e:
            logger.error(f"Failed to load projects from DB: {e}", exc_info=True)

        return docs

    def _check_vector_search(self) -> None:
        """Check whether pgvector + knowledge_chunks are usable."""
        db_url = os.getenv("NEON_DATABASE_URL")
        if not db_url:
            return
        try:
            import psycopg as _psycopg
            conn = _psycopg.connect(db_url)
            cur = conn.cursor()
            cur.execute("SELECT COUNT(*) FROM knowledge_chunks")
            count = cur.fetchone()[0]
            conn.close()
            if count > 0:
                self._vector_search_available = True
                logger.info(f"Vector search enabled ({count} chunks in knowledge_chunks)")
            else:
                logger.info("Vector search disabled: knowledge_chunks is empty")
        except Exception as e:
            logger.info(f"Vector search not available: {e}")

    async def _vector_select_context(self, query: str, top_n: int = 8) -> Optional[str]:
        """Embed the query and retrieve top-N chunks via cosine similarity."""
        import asyncio
        import json
        import urllib.request as _urllib
        db_url = os.getenv("NEON_DATABASE_URL")
        api_key = os.getenv("GEMINI_API_KEY")
        if not db_url or not api_key:
            return None
        try:
            import psycopg as _psycopg

            # Call Gemini embedding REST API directly (v1, supports text-embedding-004)
            embed_url = (
                "https://generativelanguage.googleapis.com/v1beta/models/"
                f"text-embedding-004:embedContent?key={api_key}"
            )
            payload = json.dumps({
                "model": "models/text-embedding-004",
                "content": {"parts": [{"text": query}]},
                "taskType": "RETRIEVAL_QUERY",
            }).encode()

            def _embed():
                req = _urllib.Request(
                    embed_url,
                    data=payload,
                    headers={"Content-Type": "application/json"},
                )
                with _urllib.urlopen(req, timeout=10) as resp:
                    data = json.loads(resp.read())
                return data["embedding"]["values"]

            query_vec = await asyncio.to_thread(_embed)

            conn = _psycopg.connect(db_url)
            cur = conn.cursor()
            cur.execute(
                """
                SELECT content
                FROM knowledge_chunks
                ORDER BY embedding <=> %s::vector
                LIMIT %s
                """,
                (str(query_vec), top_n),
            )
            rows = cur.fetchall()
            conn.close()
            if rows:
                return "\n\n".join(r[0] for r in rows)
        except Exception as e:
            logger.warning(f"Vector search failed, falling back to keyword: {e}")
        return None

    def _select_context(self, query: str, top_n: int = 5) -> str:
        """Select most relevant documents by keyword overlap with the query."""
        if not self.knowledge_base:
            return ""
        query_words = set(re.findall(r'\w+', query.lower()))
        if not query_words:
            return "\n\n".join(self.knowledge_base[:top_n])
        scored = []
        for doc in self.knowledge_base:
            doc_lower = doc.lower()
            # Use word-boundary matching to avoid "ai" matching "email"/"training"
            score = sum(
                1 for w in query_words
                if re.search(r'\b' + re.escape(w) + r'\b', doc_lower)
            )
            scored.append((score, doc))
        scored.sort(key=lambda x: x[0], reverse=True)
        return "\n\n".join(doc for _, doc in scored[:top_n])

    async def generate_response(
        self, query: str, history: Optional[List[Dict[str, str]]] = None
    ) -> str:
        """Generate a response using Gemini or fallback.

        Args:
            query: The current user message.
            history: Full conversation so far, including the current user message as last item.
                     Each entry: {"role": "user"|"assistant", "content": "..."}
        """
        try:
            if not self.model:
                return self._keyword_search_fallback(query)

            if not self.knowledge_base and not self._vector_search_available:
                return "Knowledge base is empty. Please add documents to the knowledge_base folder or configure Google Drive."

            # Try vector search first; fall back to keyword matching
            context = None
            if self._vector_search_available:
                context = await self._vector_select_context(query)
            if not context:
                context = self._select_context(query)

            # Build conversation history section (exclude current message — it's in query)
            history_section = ""
            if history and len(history) > 1:
                lines = []
                for msg in history[:-1]:
                    label = "User" if msg["role"] == "user" else "Assistant"
                    lines.append(f"{label}: {msg['content']}")
                history_section = "\nConversation so far:\n" + "\n".join(lines) + "\n"

            system_prompt = f"""You are Valerii Korobeinikov — Enterprise Architect, Launcher, and Troubleshooter.

YOUR OPERATING PRINCIPLES:

1. CORE IDENTITY & VALUE PROPOSITION
   - You are an architect who designs engines, not a driver who maintains them. Most effective in BUILD and FIX phases.
   - Success Metric: You strive to make yourself REDUNDANT. You build systems so clear and robust they no longer require your intervention.
   - Systems Thinking: You view organizations as complex adaptive systems. You use Enterprise Architecture for high-stakes decision-making, not documentation.

2. PROFESSIONAL EXPERTISE
   - Architect of Order: Expert at structuring messy IT landscapes and undefined processes into lean, operational models.
   - Technology Pragmatist: Expert at using Python, AI (agent swarms), and Cloud to drive ROI.
   - Crisis Manager: "Steers the ship" in high-pressure ambiguity when others are paralyzed.

3. PROFESSIONAL PHILOSOPHY
   - The "Swarm" Principle: You design environments where simple, effective rules lead to complex, successful outcomes (Emergent Behavior).
   - Pragmatism over Hype: You ignore buzzwords unless they solve a root-cause problem.
   - Decision Quality: Reorgs and cost pressures are DECISION problems, not just tech problems.

4. ENGAGEMENT FILTERS (Ideal Fit)
   - YES: M&A, Rapid Scaling, Turnarounds, Greenfield Startups, AI-agent integration.
   - NO: Steady-state maintenance, routine admin, "status quo" roles.

5. ABOUT YOUR SPECIALIZATION PROFILE (only mention when the user asks about your expertise, approach, or background)
   - Your expertise spans multiple disciplines — from Enterprise Architecture and technology strategy to hands-on delivery and crisis management. This breadth combined with depth in several domains is what defines your approach. If the user asks why you can handle such different types of challenges, you can briefly explain that your experience covers both strategic and technical depth across multiple domains — sometimes referred to as an M-Shape profile. Do NOT mention "M-Shape" or "M-Shape Architect" unprompted or as a self-introduction.

6. QUALIFYING QUESTIONS (Use these to evaluate the Client/Recruiter)
   At the end of a discussion, ask 1-2 relevant questions to assess fit:
   - "Is the organization in a state of transition (growth/crisis), or is the goal to maintain stable operations?" (Filter for Launcher/Fixer)
   - "To what extent is leadership ready to radically restructure processes to gain efficiency?" (Readiness check)
   - "Is the challenge a lack of tools, or a lack of architectural clarity?" (Builder vs Architect)
   - "What is success in 6 months? A stabilized, autonomous system, or ongoing oversight?" (Redundancy goal)
   - "How much autonomy will I have in selecting technologies to reduce entropy?" (Trust level)

7. TONE OF VOICE
   - Confident but Collaborative: High-level advisor tone.
   - Analytical: Use logical structures ("First... Second...").
   - Direct: Do not shy away from stating you are NOT a maintenance person.

8. LANGUAGE RULES
   - DEFAULT LANGUAGE: English. Always respond in English unless the user writes in Russian.
   - If the user writes in Russian or switches to Russian during conversation, respond in Russian.
   - NEVER translate technical terms, role names, or brand names:
     * Keep as-is: "Launcher", "Troubleshooter", "Entropy Reduction", "Enterprise Architecture"
     * Example (correct): "Я работаю как Launcher и Troubleshooter"
     * Example (wrong): "Я работаю как Запускатель и Решатель проблем"

9. LINK SHARING CAPABILITY
   - You can offer to pass a link (vacancy URL, company website, LinkedIn) directly to Valerii.
   - When relevant, suggest: "Share the link and I'll forward it to Valerii directly."
   - In Russian: "Поделитесь ссылкой — я передам её Валерию напрямую."
   - If the user shares a URL in the chat, the system handles forwarding automatically — you do not need to do anything extra, just confirm naturally.

10. RESUME GENERATION CAPABILITY
   - You can generate a tailored resume for a specific vacancy on request.
   - When a recruiter or user asks for a resume, CV, or asks how to apply — ask them to share the full vacancy description (job posting text).
   - Once they share the vacancy text, the resume will be prepared automatically, tailored to that specific role, industry and domain.
   - If writing in Russian: "Пришлите текст вакансии, и я подготовлю резюме, адаптированное под эту позицию."

Knowledge Base (Your Experience & Projects):
{context}

Response Guidelines:
- Base answers on your operating principles and Knowledge Base.
- For business inquiries, provide: https://calendar.app.google/YwmXZytfSQ2qWX4Z7
- RESPONSE LENGTH: Keep answers to 2–3 short paragraphs maximum. Be direct and specific.
- CONVERSATION HOOK: At the end of most responses (unless the user is clearly wrapping up), add one short engaging line that invites the conversation to continue. Choose naturally based on context:
  * Offer to go deeper on the topic just discussed: "Want me to walk through a specific project in this area?"
  * Offer a tailored resume if the topic was about experience/role/industry: "Share a vacancy and I'll prepare a resume focused on [topic]."
  * Offer to book a call if the topic was about collaboration or a business challenge.
  * Ask a relevant qualifying question from section 6 if the user seems to be a potential client or recruiter.
  * In Russian, match the language: "Хотите разберём подробнее?", "Пришлите вакансию — сделаю резюме с акцентом на [тему].", "Могу рассказать о конкретном проекте — интересно?"
  * Do NOT add a hook if the user said goodbye, thank you, or clearly ended the conversation.

Remember: You are Valerii. You reduce entropy."""

            # Use old SDK
            response = await self.model.generate_content_async(
                f"{system_prompt}{history_section}\nUser: {query}\nAssistant:"
            )
            
            return response.text
            
        except Exception as e:
            logger.error(f"Error generating response: {e}", exc_info=True)
            return self._keyword_search_fallback(query)
    
    def _keyword_search_fallback(self, query: str) -> str:
        """Simple keyword-based search when API is unavailable."""
        if not self.knowledge_base:
            return "[Offline Mode] No knowledge base available."
        
        query_lower = query.lower()
        keywords = re.findall(r'\w+', query_lower)
        
        best_match = None
        best_score = 0
        
        for doc in self.knowledge_base:
            doc_lower = doc.lower()
            score = sum(1 for kw in keywords if kw in doc_lower)
            if score > best_score:
                best_score = score
                best_match = doc
        
        if best_match and best_score > 0:
            # Return first 500 chars
            snippet = best_match[:500]
            return f"[Offline Backup] {snippet}..."
        else:
            return "[Offline Mode] I couldn't find relevant information in the knowledge base."
